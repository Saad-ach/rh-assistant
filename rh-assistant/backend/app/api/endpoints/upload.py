from fastapi import APIRouter, File, UploadFile, Depends, HTTPException
from sqlalchemy.orm import Session
import docx
import PyPDF2
import io
from app.database import get_db
from app.models import schemas, models
from app.core.config import settings # Import settings
from app.ml.embeddings import EmbeddingsGenerator
from app.services import hr_service
from app.ml.vectorizer import chroma_vectorizer
import uuid
from .chat import get_current_user
from fastapi import status

router = APIRouter()

async def extract_text_from_pdf(file: UploadFile) -> str:
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(await file.read()))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

async def extract_text_from_docx(file: UploadFile) -> str:
    document = docx.Document(io.BytesIO(await file.read()))
    text = "\n".join([paragraph.text for paragraph in document.paragraphs])
    return text

@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    category: str = "general",
    db: Session = Depends(get_db),
    current_user: schemas.User = Depends(get_current_user),
):
    # For simplicity, only admin can upload documents
    if current_user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not enough permissions")

    # Avoid reading the file twice; branch by extension and read once
    file_extension = (file.filename or "").split(".")[-1].lower()
    extracted_text = ""

    if file_extension == "pdf":
        extracted_text = await extract_text_from_pdf(file)
    elif file_extension in ["doc", "docx"]:
        extracted_text = await extract_text_from_docx(file)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF and DOCX are supported.")

    # Store in vector DB regardless of relational DB availability
    doc_id = str(uuid.uuid4())
    chroma_vectorizer.add_document(
        doc_id,
        extracted_text,
        {"source": f"upload_by_{current_user.email}", "filename": file.filename, "file_type": file_extension, "category": category}
    )

    # Optionally persist metadata in relational DB (best-effort)
    # Best effort DB write disabled in dev mode if tables missing
    try:
        hr_document = schemas.HRDocument(
            title=file.filename,
            content=extracted_text,
            source=f"upload_by_{current_user.email}",
            category=category,
            metadata={"filename": file.filename, "file_type": file_extension},
        )
        hr_service.create_hr_document(db, hr_document)
    except Exception:
        ...

    return {"filename": file.filename, "message": "Document uploaded and processed successfully", "document_id": doc_id}
